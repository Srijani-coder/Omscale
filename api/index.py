# Flask entry point
from flask import Flask, render_template, request, jsonify
from chatbot.chain import get_bot_response
import os
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from dotenv import load_dotenv
import os
import requests, io
import docx
from urllib.parse import quote_plus
import base64

load_dotenv()
SENDER_EMAIL = os.getenv("EMAIL_USER")
PASSWORD = os.getenv("EMAIL_PASS")
RAPID_KEY = os.getenv("RAPID_KEY")
RAPID_HOST = os.getenv("RAPID_HOST")

if not SENDER_EMAIL or not PASSWORD:
    raise RuntimeError("EMAIL_USER and EMAIL_PASS must be set in .env")

# Remove spaces from password only once
PASSWORD = PASSWORD.replace(" ", "")

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_DIR = os.path.join(BASE_DIR, "templates")
STATIC_DIR = os.path.join(BASE_DIR, "static")

all_articles = []  

def list_folder_files():
    url = "https://terabox-downloader-online-viewer-player-api.p.rapidapi.com/rapidapi"
    querystring = {"url": quote_plus ('https://1024terabox.com/s/1M9AMaI-1xiOHzDWtaZ-KkA')}
    headers = {
        "x-rapidapi-key": RAPID_KEY,
        "x-rapidapi-host": RAPID_HOST
    }
    response = requests.get(
        "https://terabox-downloader-online-viewer-player-api.p.rapidapi.com/rapidapi",
        headers=headers,
        params=querystring
    )
    data = response.json()
    articles = []

    for f in data.get("list", []):
        if f.get("filename", "").endswith(".docx") and "download" in f:
            # Download and load the DOCX file
            doc_bytes = requests.get(f["download"]).content
            doc = docx.Document(io.BytesIO(doc_bytes))
            # Extract title and content text
            title = doc.paragraphs[0].text if doc.paragraphs else "Untitled"
            content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            # Extract the first image (assumed to be above title)
            image_data = None
            rels = doc.part._rels

            for rel in rels:
                if "image" in rels[rel].target_ref:
                    image_part = rels[rel]._target
                    image_bytes = image_part.blob
                    image_data = base64.b64encode(image_bytes).decode("utf-8")
                    break  # Use only the first image

        # Save the article entry
        articles.append({
            "title": title,
            "content": content,
            "image_data": image_data  # base64 string of image
        })

        if len(articles) >= 3:
            break
    
    all_articles = articles
    return articles


app = Flask(__name__, template_folder=TEMPLATE_DIR, static_folder=STATIC_DIR)

@app.route('/')
def home():
    return render_template('home.html')


@app.route('/send_contact', methods=['POST'])
def send_contact():
    data = request.get_json()
    name = data.get('name')
    info = data.get('info')
    message = data.get('message')

    email_body = f"""
    New Contact Message from Website:

    Name: {name}
    Contact: {info}

    Message:
    {message}
    """

    # ✅ Use global SENDER_EMAIL & PASSWORD
    receiver_email = SENDER_EMAIL

    msg = MIMEMultipart()
    msg['From'] = SENDER_EMAIL
    msg['To'] = receiver_email
    msg['Subject'] = f"New Contact Form Submission from {name}"
    msg.attach(MIMEText(email_body, 'plain'))

    try:
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(SENDER_EMAIL, PASSWORD)
            server.sendmail(SENDER_EMAIL, receiver_email, msg.as_string())
        return jsonify({"status": "success"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})


@app.route('/services')
def services():
    services = [
        {"id": 1, "title": "Immigration Law", "description": "Expert advice for visas and residency.", "image": "immigration.png"},
        {"id": 2, "title": "Contract Law", "description": "Drafting and reviewing contracts.", "image": "contract.png"},
        {"id": 3, "title": "SOLICITOR’S PROCUREMENT AND CONTRACT MANAGEMENT", "description": "contracts procure management by solicitors.", "image": "contract_proc.png"},
        {"id": 4, "title": "CRIMINAL LAW", "description": "Defending rights with integrity.", "image": "criminal.png"},
        {"id": 5, "title": "CIVIL LAW", "description": "Protecting civil rights for all aspect ", "image": "civil.png"},
        {"id": 6, "title": "CONSTRUCTION LAW", "description": "Advise on constructing a building in legal manner", "image": "const.png"},
        {"id": 7, "title": "FAMILY LAW", "description": "Supporting families through legal challenges.", "image": "family.png"},
        {"id": 8, "title": "Wills & Estates", "description": "Planning for the future with care.", "image": "wills.png"},
        {"id": 9, "title": "Climate change & Environmental Law", "description": "Advising on sustainability and regulations.", "image": "environment.png"}
    ]
    return render_template('services.html', services=services[:9])  # limit 9 articles

@app.route('/service/<int:service_id>')
def service_article(service_id):
    articles = {
        1: {
            "title": "Immigration Law",
            "image": "immigration.png",
            "content": """
                <ul>
                    <li>Assess eligibility for Aust visa types (e.g. student, skilled, partner, work, business, protection, refugee, asylum), provide strategic advice on most appropriate visa pathway, lodged accurate visa applications, assisted with bridging visas, and maintained lawful status during visa transitions to PR and citizenship.</li>
                    <li>Stay updated with evolving migration laws and advised clients on appropriate visa pathways/eligibility.</li>
                    <li>Applied, represented and advocated for unsuccessful visa clients on complex issues (breaches, section 501 character issues) with legal arguments and evidence at the DHA, AAT (merits review), Federal Circuit & Family, Federal Court, High Court (judicial review) and Immigration Ministerial Intervention.</li>
                    <li>Helped businesses to become approved sponsors and managed employer-sponsored visa processes.</li>
                </ul>
            """
        },
        2: {
            "title": "Contract Law",
            "image": "contract.png",
            "content": """
                <ul>
                    <li>Drafted clear, precise, and legally enforceable contracts tailored to specific needs (e.g. service agreements, supplier/vendor contracts) while ensuring compliance with Australian Competition and Consumer Act, Corporations Act, Fair Work Act.</li>
                    <li>Reviewed existing contracts for legal risks, ambiguities, compliance and advised clients on their rights, obligations and potential legal consequences of contract terms.</li>
                    <li>Assisted clients in negotiating favorable terms and conditions, while protecting client interests and promoting mutually agreeable outcomes</li>
                    <li>Handled breaches of contract/disputes through negotiation, mediation, arbitration, or litigation, and represented clients in court/tribunal proceedings on contract disputes.</li>
                    <li>Advised on the legal grounds for terminating contracts (e.g. breach, frustration, termination clauses) and enforced contractual rights and pursued remedies such as damages or specific performance.</li>
                    <li>Advised on international contracts, including jurisdiction, choice of law, dispute resolution clauses and international trade laws/treaties (e.g. CISG).</li>
                </ul>
            """
        },
        3: {
            "title": "SOLICITOR’S PROCUREMENT AND CONTRACT MANAGEMENT",
            "image": "contract_proc.png",
            "content": """
                <ul>
                    Advised and complied with Commonwealth Procurement Rules (CPRs), PGPA Act, ensuring value-for-money in procurement processes. Drafted, reviewed and negotiated a range of enforceable commercial contracts, e.g. procurement & supply contracts, SLAs, ICT, PPP contracts. Advised and supported clients on the legal structure of EOI, RFT, tender evaluation, probity, ethical standards, legal and commercial risks and mitigation strategies (e.g. indemnities, limitations of liability), negotiations with vendors/suppliers/contractors, contract breaches, disputes resolution (negotiation, mediation, arbitration, litigation, terminations). Liaised with government regulators (e.g. Australian National Anti-Corruption Commission) to deter fraud in Government Procurement and Contract. Trained, developed, supervised, coached and mentored APS employees in procurement and contract laws.
                </ul>
            """
        },
        4: {
            "title": "CRIMINAL LAW",
            "image": "criminal.png",
            "content": """
                <ul>
                    <li>Advised clients on their legal rights, obligations, and potential charges. Explained criminal charges, court processes and likely outcomes. Assisted in police interviews and investigations to protect client rights.</li>
                    <li>Represented clients in all criminal court levels: Local, District, LEC, Supreme Court and appellate courts. Appeared in bail, plea hearings, trials, and sentencing. Advocated on behalf of the accused to ensure a fair trial and just outcome.</li>
                    <li>Analysed evidence, police briefs, and witness statements. Investigated facts, interviewed witnesses and engaged forensic, psychological experts if needed and developed a defence strategy or negotiate charges where appropriate.</li>
                    <li>Negotiated with the prosecution to reduce or drop charges (plea bargaining). Seek alternative sentencing options or diversionary programs (e.g., drug court).</li>
                    <li>Prepared and filed legal documents: bail applications, submissions, appeals and court motions. Drafted character references, sentencing submissions and briefs for trial.</li>
                    <li> Advised on the merits of appeals and represented clients in appellate courts for sentence appeals or convictions. Assisted with parole applications or miscarriage of justice claims.</li>
                </ul>
            """
        },
        5: {
            "title": "CIVIL LAW",
            "image": "civil.png",
            "content": """
                <ul>
                    <li>Advised clients on their legal rights, responsibilities in civil matters and identified legal issues, assessed risks and outline possible outcomes and remedies.</li>
                    <li>Drafted pleadings, claims, defence documents, affidavits and reviewed and drafted contracts, deeds, agreements and settlement offers.</li>
                    <li>Attempted to resolve disputes through negotiation, mediation, arbitration before proceeding to court. 4. Represented clients in ADR forums such as: Mediation, Conciliation and Arbitration.</li>
                    <li>Develop litigation strategies based on the strength of the case, evidence, and client goals.</li>
                    <li>Assisted clients in enforcing court orders (e.g. garnishment, property seizure, or bankruptcy proceedings</li>
                </ul>
            """
        },
        6: {
            "title": "CONSTRUCTION LAW",
            "image": "const.png",
            "content": """
                <ul>
                    <li>Drafted, reviewed construction contracts (e.g. AS 4000, AS 2124 contracts), advised on subcontractor agreements, consultancy agreements, and procurement terms.</li>
                    <li>Advised on compliance with: Building Code of Australia (BCA), WHS Act 2011 and ensured environmental, planning, and development approvals</li>
                    <li>Identified and mitigated legal risks during project lifecycle including Public liability, Professional indemnity insurance.</li>
                    <li>Implemented procedures for representing clients in civil proceedings before courts and tribunals such as: NCAT.</li>
                    <li>Developed litigation strategies based on the case, evidence and client goals.</li>
                </ul>
            """
        },
        7: {
            "title": "FAMILY LAW",
            "image": "family.png",
            "content": """
                <ul>
                    <li>Advised clients on legal requirements for divorce. Assisted in filing divorce applications in the Federal Circuit and Family Court of Australia. Handled legal aspects of de facto relationship breakdowns.</li>
                    <li>Helped parents agree on parenting plans, orders (living arrangements, decision-making, communication). Represented clients in custody disputes and advocated for the child’s best interests. Addressed issues like relocation, child abduction and parenting time schedules.</li>
                    <li>Advised clients on division of assets, debts and superannuation following separation. Negotiated and formalised property settlement agreements. Applied for consent orders or represented clients in court property disputes.</li>
                    <li>Advised on eligibility and applied for spousal maintenance. Negotiated ongoing/lump-sum support payments. Assisted with Child Support Agency interactions or private agreements.</li>
                    <li>Drafted pre-nuptial, post-nuptial, and cohabitation agreements and their enforcement.</li>
                    <li>Assisted clients with Domestic Violence Orders (DVOs) or Apprehended Violence Orders (AVOs). Represented clients in court and ensured safety and protection under the law.</li>
                    <li>Encouraged clients to resolve matters through mediation, arbitration or collaborative law. Represented clients in Family Dispute Resolution (FDR) sessions.</li>
                    <li>Represented clients in family law courts during trials, hearings, appeals, drafing legal documentation, affidavits, and submissions.</li>
                </ul>
            """
        },
        8: {
            "title": "Wills & Estates",
            "image": "wills.png",
            "content": """
                <ul>
                    <li>Creating valid and legally enforceable wills that reflect the client’s intentions.</li>
                    <li>Setting up trusts in wills for beneficiaries, often to protect vulnerable family members or for tax purposes.</li>
                    <li>Preparing enduring powers of attorney, guardianship documents, and advance care directives.</li>
                    <li>Advising how to structure personal and business assets to achieve the best estate planning outcomes.</li>
                    <li>Advising on estate administration and disputes</li>
                <ul>
            """
        },
        9: {
            "title": "Climate change & Environmental Law",
            "image": "environment.png",
            "content": """
                <ul>
                    <li>Advising clients on obligations under emissions reporting schemes (e.g., NGER, Safeguard Mechanism), climate-related disclosure requirements, renewable energy project development,  climate change adaptation and mitigation planning</li>
                    <li>Representing clients facing prosecution by regulators (e.g., NSW EPA, local councils, Commonwealth Clean Energy Regulator) for alleged environmental offences</li>
                </ul>
            """
        },
        

    }
    article = articles.get(service_id)
    if not article:
        return render_template('404.html'), 404
    return render_template('service_article.html', article=article)
    


@app.route("/dashboard")
def dashboard():
    articles = list_folder_files()
    return render_template("dashboard.html", articles=articles)

@app.route("/article/<int:article_id>", methods=["GET", "POST"])
def blog_article(article_id):
    article = all_articles[article_id]

    # Handle comment form submission
    if request.method == "POST":
        comment = request.form.get("comment")
        if comment:
            article.setdefault("comments", []).append(comment)

    return render_template("blog_article.html",
                           title=article["title"],
                           content=article["content"],
                           image_data=article["image_data"],
                           comments=article.get("comments", []))

@app.route('/chatbot')
def chatbot():
    return render_template('chatbot.html')

@app.route('/get_response', methods=['POST'])
def get_response():
    user_query = request.json.get("query")
    response = get_bot_response(user_query)
    return jsonify({"response": response})

if __name__ == '__main__':
    app.run(debug=True)